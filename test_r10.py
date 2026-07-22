"""Test R10: DOCX + XLSX support"""
import sys, logging
sys.path.insert(0, ".")
logging.basicConfig(level=logging.WARNING)

from src.document_processor import DocumentProcessor
proc = DocumentProcessor(chunk_size=512, chunk_overlap=64)

print("=== Test R10: DOCX + XLSX ===")

# Tao file DOCX test
print("\n[1] Tao file DOCX test...")
try:
    from docx import Document as DocxDoc
    d = DocxDoc()
    d.add_heading("RAG Overview", level=1)
    d.add_paragraph("Retrieval-Augmented Generation (RAG) combines retrieval and generation.")
    d.add_heading("Methods", level=2)
    d.add_paragraph("Dense retrieval uses vector embeddings for semantic search.")
    d.add_paragraph("BM25 is a classic keyword-based retrieval method.")
    t = d.add_table(rows=2, cols=3)
    t.rows[0].cells[0].text = "Method"
    t.rows[0].cells[1].text = "Type"
    t.rows[0].cells[2].text = "Speed"
    t.rows[1].cells[0].text = "BM25"
    t.rows[1].cells[1].text = "Keyword"
    t.rows[1].cells[2].text = "Fast"
    d.save("test_sample.docx")
    print("   DOCX file created")

    docs = proc.process("test_sample.docx")
    print(f"   DOCX chunks: {len(docs)}")
    print(f"   Sample: {docs[0].page_content[:100]}")
    print(f"   Processor: {docs[0].metadata['processor']}")
    print("   DOCX: PASS")
except Exception as e:
    print(f"   DOCX: FAIL - {e}")

# Tao file XLSX test
print("\n[2] Tao file XLSX test...")
try:
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "RAG Metrics"
    ws.append(["Model", "Faithfulness", "Relevancy", "Precision"])
    ws.append(["RAG-base", 0.72, 0.68, 0.75])
    ws.append(["RAG-hybrid", 0.85, 0.81, 0.88])
    ws.append(["RAG-reranker", 0.91, 0.87, 0.93])
    wb.save("test_sample.xlsx")
    print("   XLSX file created")

    docs = proc.process("test_sample.xlsx")
    print(f"   XLSX chunks: {len(docs)}")
    print(f"   Sample: {docs[0].page_content[:150]}")
    print(f"   Sheet: {docs[0].metadata.get('sheet_name')}")
    print(f"   Headers: {docs[0].metadata.get('headers')}")
    print("   XLSX: PASS")
except Exception as e:
    print(f"   XLSX: FAIL - {e}")

# Cleanup
import os
for f in ["test_sample.docx", "test_sample.xlsx"]:
    if os.path.exists(f):
        os.remove(f)

print("\n=== R10 COMPLETE ===")
