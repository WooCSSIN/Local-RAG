"""
document_processor.py — Xử lý tài liệu đa định dạng
Hỗ trợ: PDF, Word (.docx), Excel (.xlsx), HTML, Markdown, TXT
R10: Thêm DOCX (python-docx fallback) + XLSX (openpyxl fallback)
Dùng Docling với fallback khi không có.
"""
import logging
from pathlib import Path

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

logger = logging.getLogger(__name__)

# Định dạng Docling hỗ trợ
DOCLING_SUPPORTED = {".pdf", ".docx", ".xlsx", ".html", ".htm", ".md", ".markdown"}
TEXT_FORMATS = {".txt"}
# R10: thêm .docx và .xlsx vào supported list
ALL_SUPPORTED = DOCLING_SUPPORTED | TEXT_FORMATS


class DocumentProcessor:
    """
    Xử lý tài liệu đa định dạng, trả về list[Document] sẵn để index.
    Ưu tiên Docling nếu có. Fallback riêng cho từng định dạng.
    """

    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 64):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self._docling_available = self._check_docling()
        self._splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            add_start_index=True,
            separators=["\n\n", "\n", "(?<=[.?!])", " ", ""],
        )

    def _check_docling(self) -> bool:
        try:
            from docling.document_converter import DocumentConverter  # noqa
            return True
        except ImportError:
            logger.warning("Docling chưa được cài. Dùng fallback cho từng định dạng.")
            return False

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def process(self, file_path: str) -> list[Document]:
        """
        Xử lý file và trả về list Document đã chunk.
        Hỗ trợ: PDF, DOCX, XLSX, HTML, MD, TXT
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File không tồn tại: {file_path}")

        suffix = path.suffix.lower()
        logger.info(f"Xử lý: {path.name} ({suffix})")

        # Dùng Docling nếu có và định dạng được hỗ trợ
        if suffix in DOCLING_SUPPORTED and self._docling_available:
            docs = self._process_with_docling(str(path))

        # Fallback theo từng định dạng
        elif suffix == ".pdf":
            docs = self._process_pdf_fallback(str(path))
        elif suffix == ".docx":
            docs = self._process_docx(str(path))          # R10
        elif suffix == ".xlsx":
            docs = self._process_xlsx(str(path))          # R10
        elif suffix in {".html", ".htm"}:
            docs = self._process_html(str(path))
        elif suffix in {".md", ".markdown", ".txt"}:
            docs = self._process_text(str(path))
        else:
            raise ValueError(
                f"Định dạng '{suffix}' chưa hỗ trợ.\n"
                f"Hỗ trợ: {', '.join(sorted(ALL_SUPPORTED))}"
            )

        logger.info(f"✅ {len(docs)} chunks từ {path.name}")
        return docs

    def process_multiple(self, file_paths: list[str]) -> list[Document]:
        all_docs = []
        for fp in file_paths:
            try:
                all_docs.extend(self.process(fp))
            except Exception as e:
                logger.error(f"Lỗi xử lý {fp}: {e}")
        return all_docs

    # ------------------------------------------------------------------
    # Docling (best quality)
    # ------------------------------------------------------------------

    def _process_with_docling(self, file_path: str) -> list[Document]:
        from docling.document_converter import DocumentConverter
        converter = DocumentConverter()
        result = converter.convert(file_path)
        dl_doc = result.document
        try:
            from docling.chunking import HybridChunker
            chunker = HybridChunker(max_tokens=self.chunk_size, merge_peers=True)
            chunks = list(chunker.chunk(dl_doc=dl_doc))
            docs = []
            for chunk in chunks:
                meta = {}
                try:
                    meta = chunk.meta.export_json_dict()
                except Exception:
                    pass
                docs.append(Document(
                    page_content=chunk.text,
                    metadata={
                        "source": file_path,
                        "filename": Path(file_path).name,
                        "page": meta.get("page_no"),
                        "heading": (meta.get("headings") or [None])[0],
                        "doc_type": Path(file_path).suffix,
                        "processor": "docling",
                    }
                ))
            return docs
        except Exception:
            md_text = dl_doc.export_to_markdown()
            return self._split_text(md_text, file_path, processor="docling-md")

    # ------------------------------------------------------------------
    # PDF fallback
    # ------------------------------------------------------------------

    def _process_pdf_fallback(self, file_path: str) -> list[Document]:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        pages_text = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                pages_text.append(f"[Trang {i+1}]\n{text}")
        return self._split_text("\n\n".join(pages_text), file_path, processor="pypdf")

    # ------------------------------------------------------------------
    # R10: DOCX support
    # ------------------------------------------------------------------

    def _process_docx(self, file_path: str) -> list[Document]:
        """
        Xử lý Word .docx dùng python-docx.
        Giữ cấu trúc heading → paragraph.
        """
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError(
                "python-docx chưa cài. Chạy: pip install python-docx"
            )

        doc = DocxDocument(file_path)
        lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            # Giữ heading dưới dạng markdown
            style = para.style.name if para.style else ""
            if "Heading" in style:
                level = style.replace("Heading ", "").strip()
                try:
                    hashes = "#" * int(level)
                    lines.append(f"{hashes} {text}")
                except ValueError:
                    lines.append(f"## {text}")
            else:
                lines.append(text)

        # Xử lý tables
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                lines.append("\n".join(rows))

        full_text = "\n\n".join(lines)
        if not full_text.strip():
            raise ValueError(f"File DOCX rỗng hoặc không có text: {file_path}")

        return self._split_text(full_text, file_path, processor="python-docx")

    # ------------------------------------------------------------------
    # R10: XLSX support
    # ------------------------------------------------------------------

    def _process_xlsx(self, file_path: str) -> list[Document]:
        """
        Xử lý Excel .xlsx dùng openpyxl.
        Mỗi sheet → text với header context.
        """
        try:
            import openpyxl
        except ImportError:
            raise ImportError("openpyxl chưa cài. Chạy: pip install openpyxl")

        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        all_docs = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                continue

            # Hàng đầu là header
            headers = [str(h) if h is not None else f"Cột_{i}" for i, h in enumerate(rows[0])]

            # Mỗi hàng data → text có context header (R10 requirement)
            text_rows = [f"Sheet: {sheet_name} | Cột: {', '.join(headers)}"]
            for row in rows[1:]:
                cells = [
                    f"{headers[i]}: {str(v)}"
                    for i, v in enumerate(row)
                    if v is not None and str(v).strip()
                ]
                if cells:
                    text_rows.append(" | ".join(cells))

            sheet_text = "\n".join(text_rows)
            if sheet_text.strip():
                sheet_docs = self._split_text(
                    sheet_text,
                    file_path,
                    processor="openpyxl",
                )
                # Thêm metadata sheet
                for d in sheet_docs:
                    d.metadata["sheet_name"] = sheet_name
                    d.metadata["headers"] = headers
                all_docs.extend(sheet_docs)

        wb.close()
        if not all_docs:
            raise ValueError(f"File XLSX rỗng hoặc không có dữ liệu: {file_path}")
        return all_docs

    # ------------------------------------------------------------------
    # HTML + Text fallback
    # ------------------------------------------------------------------

    def _process_html(self, file_path: str) -> list[Document]:
        try:
            from bs4 import BeautifulSoup
            with open(file_path, encoding="utf-8", errors="replace") as f:
                soup = BeautifulSoup(f.read(), "html.parser")
            text = soup.get_text(separator="\n")
        except ImportError:
            with open(file_path, encoding="utf-8", errors="replace") as f:
                text = f.read()
        return self._split_text(text, file_path, processor="html")

    def _process_text(self, file_path: str) -> list[Document]:
        with open(file_path, encoding="utf-8", errors="replace") as f:
            text = f.read()
        return self._split_text(text, file_path, processor="text")

    # ------------------------------------------------------------------
    # Shared splitter
    # ------------------------------------------------------------------

    def _split_text(self, text: str, source: str, processor: str = "default") -> list[Document]:
        base_doc = Document(
            page_content=text,
            metadata={
                "source": source,
                "filename": Path(source).name,
                "doc_type": Path(source).suffix,
                "processor": processor,
            }
        )
        return self._splitter.split_documents([base_doc])


# ------------------------------------------------------------------
# Quick test
# ------------------------------------------------------------------
if __name__ == "__main__":
    import sys
    logging.basicConfig(level=logging.INFO)
    if len(sys.argv) < 2:
        print("Usage: python document_processor.py <file_path>")
        sys.exit(1)
    proc = DocumentProcessor(chunk_size=512, chunk_overlap=64)
    docs = proc.process(sys.argv[1])
    print(f"\n{len(docs)} chunks")
    print(f"Chunk 1: {docs[0].page_content[:300]}")
    print(f"Metadata: {docs[0].metadata}")
